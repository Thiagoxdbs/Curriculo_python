from docx import Document
from docx.shared import Pt

# Cria um novo documento Word
doc = Document()

# Adiciona o cabeçalho
doc.add_heading('Thiago Alexsandro Bianchi de Souza', 0)

# Adiciona informações de contato
doc.add_paragraph('(11) 98812-4773')
doc.add_paragraph('Thiago-hiago201@hotmail.com')

# Adiciona o objetivo
doc.add_heading('Objetivo', level=1)
doc.add_paragraph(
    "Contribuir para a sustentação de ambientes com foco em MongoDB, aplicando meus conhecimentos "
    "em instalação, configuração, suporte e migração de bancos de dados, visando a eficiência e a performance do sistema."
)

# Adiciona a experiência profissional
doc.add_heading('Experiência Profissional', level=1)
experiencias = [
    ("Analista Computacional - Tecnocomp LTDA", "16/05/2023 - Atual",
     "- Automação de processos utilizando Python (bibliotecas Selenium, Pandas, Pynput)\n"
     "- Manutenção e suporte de VMs e servidores\n"
     "- Manutenção de hardware e software\n"
     "- Gestão de sistemas Windows e pacotes Office"),

    ("Estágio - Desenvolvedor - Nerus", "17/06/2021 - 22/09/2022",
     "- Automação de processos em Excel e Google Sheets\n"
     "- Desenvolvimento de automações web com Python\n"
     "- Uso do Power BI para visualização de dados"),

    ("Analista de Dados - Luana Tavares", "21/12/2021 - 10/07/2022",
     "- Desenvolvimento de planos estratégicos em redes com Excel e Google Data Studio\n"
     "- Automação web utilizando Python"),

    ("Estagiário de Help Desk - Câmara Municipal dos Vereadores", "24/09/2020 - 31/12/2020",
     "- Suporte técnico N1 e N2"),

    ("Estagiário de Help Desk - Liberty Parking", "02/2019 - 12/2019",
     "- Suporte técnico N1 e N2"),
]

for title, date, desc in experiencias:
    doc.add_heading(title, level=2)
    doc.add_paragraph(f'Data: {date}')
    doc.add_paragraph(desc)

# Adiciona habilidades técnicas
doc.add_heading('Habilidades Técnicas', level=1)
habilidades = (
    "- Bancos de Dados: MongoDB, MySQL, PostgreSQL\n"
    "- Sistemas Operacionais: Red Hat Enterprise Linux (RHEL), Windows Server\n"
    "- Cloud e Ferramentas: Google Cloud, Mongo Atlas\n"
    "- Automação e Scripting: Python, Shell Script, Ansible, Terraform\n"
    "- Desenvolvimento Web: HTML, JavaScript\n"
    "- Outros: Power BI, Pacote Office, Excel Avançado, Google Analytics"
)
doc.add_paragraph(habilidades)

# Adiciona formação acadêmica
doc.add_heading('Formação Acadêmica', level=1)
doc.add_paragraph("- Engenharia da Computação (Conclusão em 06/2024)")

# Adiciona certificações e cursos
doc.add_heading('Certificações e Cursos', level=1)
cursos = (
    "- Python: \n"
    "  - Mundo 1, 2, 3 Python 3 – Curso em Vídeo - 120H\n"
    "  - Curso de Python 3 do Básico Ao Avançado - Udemy - 141H\n"
    "  - Formação Python Developer – DIO – 64H\n"
    "- Ciência de Dados: Potência Tech powered by iFood | Ciência de Dados – DIO – 80H (em andamento)\n"
    "- Database: Database Experience - DIO - 54H\n"
    "- Programação: Potência Tech iFood - Programação do Zero - DIO - 68H\n"
    "- Outros: Web Design, Desenvolvimento Web com PHP, Banco de Dados MySQL, Contabilidade Tributária"
)
doc.add_paragraph(cursos)

# Adiciona idiomas
doc.add_heading('Idiomas', level=1)
doc.add_paragraph("- Inglês: Intermediário")

# Salva o arquivo Word
doc.save('Curriculo_Thiago_Alexsandro_Bianchi_de_Souza.docx')

print("Arquivo Word gerado com sucesso!")
